import sys
import os
import json
import re
import tempfile
import subprocess
import shutil
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

# --- Dependency Check ---
# Attempt imports and provide specific error messages if libraries are missing.
missing_libs = []
try:
    from PyQt5.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QPushButton, QLabel, QListWidget, QFileDialog, QMessageBox,
        QListWidgetItem, QStyleFactory
    )
    from PyQt5.QtCore import Qt
except ImportError:
    missing_libs.append("PyQt5")

try:
    import docx
    import docx2txt
except ImportError:
    missing_libs.append("python-docx, docx2txt")

try:
    from PIL import Image
except ImportError:
    missing_libs.append("Pillow")

try:
    import pytesseract
    from pytesseract import Output
    # Check if Tesseract executable is available (optional but recommended)
    try:
        pytesseract.get_tesseract_version()
    except pytesseract.TesseractNotFoundError:
        print("WARNING: Tesseract OCR engine not found or not in PATH. "
              "Image text extraction will fail.")
except ImportError:
    missing_libs.append("pytesseract")

try:
    import networkx as nx
except ImportError:
    missing_libs.append("networkx")

try:
    import fitz  # PyMuPDF
except ImportError:
    missing_libs.append("PyMuPDF")

# --- NEW: Check for pdfplumber ---
try:
    import pdfplumber
except ImportError:
    missing_libs.append("pdfplumber")
# --- REMOVED: Camelot check and Ghostscript check ---


if missing_libs:
    error_message = (
        "Critical Error: Missing required libraries.\n\n"
        "Please install the following libraries:\n"
        f"- {', '.join(missing_libs)}\n\n"
        "Example using pip:\n"
        "pip install PyQt5 python-docx docx2txt Pillow pytesseract networkx PyMuPDF pdfplumber\n\n"
        "You also need to install the Tesseract OCR engine separately "
        "and ensure it's in your system's PATH."
        # Removed Ghostscript requirement mention
    )
    print(error_message)
    # Attempt to show a graphical message box if PyQt is partially available
    try:
        app = QApplication([])
        QMessageBox.critical(None, "Dependency Error", error_message)
    except Exception:
        pass # If even QApplication fails, printing to console is the only option
    sys.exit(1)


# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


# --- Main Application Class ---
class XDoc2JSONApp(QMainWindow):
    """
    Main application window for the Document Content Extractor.

    Allows users to select .txt, .docx, and .pdf files, extracts structured
    content (text, tables, image OCR data, basic graph representations),
    and saves the output as a structured JSON or JSONL file.
    Uses PyMuPDF and pdfplumber for PDF processing.
    """
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Document Content Extractor by Usman")
        self.resize(800, 600)
        self.selected_files: List[Path] = []

        QApplication.setStyle(QStyleFactory.create('Fusion'))
        self._init_ui()

    def _init_ui(self) -> None:
        """Initializes the user interface components and layout."""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        # --- File Selection Section ---
        file_select_layout = QHBoxLayout()
        self.select_files_btn = QPushButton("Select Files")
        self.select_files_btn.setToolTip("Click to select one or more document files.")
        self.select_files_btn.clicked.connect(self.select_files)
        file_select_layout.addWidget(self.select_files_btn)
        main_layout.addLayout(file_select_layout)

        self.instruction_label = QLabel("Select .txt or .docx or .PDF files")
        self.instruction_label.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(self.instruction_label)

        # --- Selected Files List ---
        self.file_list = QListWidget()
        self.file_list.setStyleSheet("QListWidget { border: 1px solid gray; }")
        self.file_list.setToolTip("List of files selected for processing.")
        main_layout.addWidget(self.file_list)

        self.remove_file_btn = QPushButton("Remove Selected File")
        self.remove_file_btn.setToolTip("Remove the highlighted file from the list.")
        self.remove_file_btn.clicked.connect(self.remove_selected_file)
        main_layout.addWidget(self.remove_file_btn)

        # --- Action Buttons ---
        action_btn_layout = QHBoxLayout()
        self.extract_json_btn = QPushButton("Extract to JSON")
        self.extract_json_btn.setToolTip("Process selected files and save structured output as a single JSON file.")
        self.extract_json_btn.clicked.connect(lambda: self.run_extraction("json"))
        action_btn_layout.addWidget(self.extract_json_btn)

        self.extract_jsonl_btn = QPushButton("Extract to JSONL")
        self.extract_jsonl_btn.setToolTip("Process selected files and save structured output as a JSONL file (one JSON object per line).")
        self.extract_jsonl_btn.clicked.connect(lambda: self.run_extraction("jsonl"))
        action_btn_layout.addWidget(self.extract_jsonl_btn)
        main_layout.addLayout(action_btn_layout)

        # --- Status Label ---
        self.status_label = QLabel("Awaiting user action")
        self.status_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        self.status_label.setAlignment(Qt.AlignCenter)
        self.status_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        main_layout.addWidget(self.status_label)

        main_layout.addStretch(1)

    # -------------------------------
    # File selection handlers
    # -------------------------------
    def select_files(self) -> None:
        """Opens a file dialog to select multiple document files."""
        # Use the current time to inform the user about potential outdatedness
        current_time_str = f"{self.current_time_str()}" # Assuming method exists or add it
        print(f"Current time: {current_time_str}") # Inform in console

        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Select Document Files",
            "",
            "Documents (*.txt *.docx *.pdf)"
        )
        if files:
            newly_added_count = 0
            for file_str in files:
                file_path = Path(file_str)
                if file_path not in self.selected_files:
                    self.selected_files.append(file_path)
                    item = QListWidgetItem(file_path.name)
                    item.setData(Qt.UserRole, file_path)
                    self.file_list.addItem(item)
                    newly_added_count += 1
            if newly_added_count > 0:
                 logging.info(f"Added {newly_added_count} file(s) to the list.")
            self.instruction_label.setVisible(self.file_list.count() == 0)

    def remove_selected_file(self) -> None:
        """Removes the currently selected file from the list and internal tracking."""
        selected_items = self.file_list.selectedItems()
        if not selected_items:
            logging.warning("Remove button clicked but no file selected.")
            return

        for item in selected_items:
            file_path_to_remove = item.data(Qt.UserRole)
            if file_path_to_remove in self.selected_files:
                self.selected_files.remove(file_path_to_remove)
                logging.info(f"Removed file: {file_path_to_remove.name}")
            else:
                 logging.warning(f"File {file_path_to_remove.name} was in list but not in internal tracking.")
            self.file_list.takeItem(self.file_list.row(item))

        self.instruction_label.setVisible(self.file_list.count() == 0)

    # -------------------------------
    # Main Extraction Process
    # -------------------------------
    def run_extraction(self, output_format: str) -> None:
        """
        Initiates the content extraction process for selected files.

        Args:
            output_format: The desired output format ('json' or 'jsonl').
        """
        if not self.selected_files:
            QMessageBox.warning(
                self, "No Files Selected", "Please select one or more document(s) for processing."
            )
            return

        file_count = len(self.selected_files)
        file_plural = "file" if file_count == 1 else "files"
        reply = QMessageBox.question(
            self,
            "Confirm Extraction",
            f"Do you want to start the extraction process for the {file_count} selected {file_plural}?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            self.status_label.setText("Extraction cancelled by user.")
            logging.info("User cancelled the extraction process.")
            return

        self.status_label.setText("Extraction is in process; please wait!")
        self.set_ui_enabled(False)
        QApplication.processEvents()

        try:
            with tempfile.TemporaryDirectory() as temp_dir_str:
                temp_dir = Path(temp_dir_str)
                logging.info(f"Created temporary directory: {temp_dir}")

                all_results: Dict[str, Dict[str, Any]] = {}
                has_errors = False

                for file_path in self.selected_files:
                    logging.info(f"Processing file: {file_path.name}")
                    try:
                        ext = file_path.suffix.lower()
                        result: Optional[Dict[str, Any]] = None
                        file_temp_dir = temp_dir / f"processing_{file_path.stem}"
                        file_temp_dir.mkdir(exist_ok=True)

                        if ext == ".txt":
                            result = self._process_txt(file_path)
                        elif ext == ".docx":
                            result = self._process_docx(file_path, file_temp_dir)
                        elif ext == ".pdf":
                            # Now uses pdfplumber for tables
                            result = self._process_pdf(file_path, file_temp_dir)
                        else:
                            logging.warning(f"Skipping unsupported file type: {file_path.name}")
                            continue

                        if result:
                            all_results[file_path.name] = result

                    except Exception as file_error:
                        has_errors = True
                        error_msg = f"Failed to process file '{file_path.name}': {file_error}"
                        logging.error(error_msg, exc_info=True)
                        all_results[file_path.name] = {"error": error_msg}

                if not all_results:
                     raise ValueError("No files were successfully processed or selected.")

                self._save_output(all_results, output_format, has_errors)

        except Exception as e:
            self.status_label.setText("Error during the extraction process.")
            error_detail = f"An unexpected error occurred: {str(e)}"
            logging.error(error_detail, exc_info=True)
            QMessageBox.critical(self, "Extraction Error", error_detail)

        finally:
            self.set_ui_enabled(True)

    def set_ui_enabled(self, enabled: bool) -> None:
        """Enable or disable key UI elements during processing."""
        self.select_files_btn.setEnabled(enabled)
        self.file_list.setEnabled(enabled)
        self.remove_file_btn.setEnabled(enabled)
        self.extract_json_btn.setEnabled(enabled)
        self.extract_jsonl_btn.setEnabled(enabled)

    def _save_output(self, results: Dict[str, Any], format: str, has_errors: bool) -> None:
        """Handles saving the processed results to a file."""
        if format == "json":
            file_filter = "JSON Files (*.json)"
            default_suffix = ".json"
        elif format == "jsonl":
            file_filter = "JSON Lines Files (*.jsonl)"
            default_suffix = ".jsonl"
        else:
            raise ValueError(f"Unsupported output format: {format}")

        save_path_str, _ = QFileDialog.getSaveFileName(
            self, "Save Extracted Output", f"extracted_output{default_suffix}", file_filter
        )

        if not save_path_str:
            self.status_label.setText("Save operation cancelled by user.")
            logging.info("User cancelled the save operation.")
            return

        save_path = Path(save_path_str)

        try:
            with open(save_path, "w", encoding="utf-8") as f:
                if format == "json":
                    json.dump(results, f, indent=4, ensure_ascii=False)
                elif format == "jsonl":
                    for file_name, data in results.items():
                        record = {"source_file": file_name, "extracted_data": data}
                        json.dump(record, f, ensure_ascii=False)
                        f.write("\n")

            logging.info(f"Output successfully saved to: {save_path}")

            if has_errors:
                 self.status_label.setText("Extraction completed with some errors (check log/output file).")
                 QMessageBox.warning(self, "Extraction Complete",
                                    f"Extraction finished, but errors occurred for one or more files.\n"
                                    f"Output saved to: {save_path}\nPlease check the output file and logs for details.")
            else:
                self.status_label.setText("Extraction completed successfully.")
                QMessageBox.information(self, "Extraction Complete",
                                        f"Extraction finished successfully.\nOutput saved to: {save_path}")

            self._open_folder(save_path.parent)

        except IOError as e:
            self.status_label.setText("Error saving the output file.")
            error_detail = f"Failed to write output file to '{save_path}': {e}"
            logging.error(error_detail, exc_info=True)
            QMessageBox.critical(self, "File Save Error", error_detail)
        except Exception as e:
            self.status_label.setText("Error during file saving.")
            error_detail = f"An unexpected error occurred while saving: {e}"
            logging.error(error_detail, exc_info=True)
            QMessageBox.critical(self, "File Save Error", error_detail)

    # -------------------------------
    # Processing methods for each file type
    # -------------------------------
    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Processes a .txt file, extracting its content."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return {
                "content_type": "text",
                "text_content": content,
                "tables": [],
                "images": {}
            }
        except Exception as e:
            logging.error(f"Error reading TXT file {file_path.name}: {e}")
            raise

    def _process_docx(self, file_path: Path, temp_dir: Path) -> Dict[str, Any]:
        """
        Processes a .docx file, extracting text, tables, and images (with OCR).
        """
        # (This method remains unchanged from the previous version)
        results: Dict[str, Any] = {
            "content_type": "docx",
            "text_content": "",
            "tables": [],
            "images": {}
        }
        image_paths: List[Tuple[str, Path]] = []

        try:
            # 1. Extract Text
            results["text_content"] = docx2txt.process(file_path)
            logging.info(f"Extracted text from {file_path.name}")
        except Exception as e:
            logging.warning(f"Could not extract text using docx2txt for {file_path.name}: {e}. Falling back.")
            try:
                doc = docx.Document(file_path)
                results["text_content"] = "\n".join([p.text for p in doc.paragraphs])
            except Exception as e_docx:
                 logging.error(f"Could not extract any text from {file_path.name}: {e_docx}")
                 results["text_content"] = "[Text extraction failed]"

        try:
            # 2. Extract Tables
            doc = docx.Document(file_path)
            for i, tbl in enumerate(doc.tables):
                table_data = []
                try:
                    for row in tbl.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    results["tables"].append({"id": f"table_{i+1}", "data": table_data})
                except Exception as e_tbl:
                     logging.warning(f"Could not fully process table {i+1} in {file_path.name}: {e_tbl}")
                     results["tables"].append({"id": f"table_{i+1}", "error": str(e_tbl)})
            logging.info(f"Extracted {len(results['tables'])} table(s) from {file_path.name}")

            # 3. Extract Images
            image_dir = temp_dir / "images"
            image_dir.mkdir(exist_ok=True)
            img_index = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        if rel.is_external: continue
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        img_ext = Path(image_part.partname).suffix.lower() or '.png'
                        if img_ext not in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']: img_ext = '.png'
                        img_filename = f"image_{img_index+1}{img_ext}"
                        img_save_path = image_dir / img_filename
                        with open(img_save_path, "wb") as f: f.write(image_bytes)
                        image_paths.append((rel.target_ref, img_save_path))
                        img_index += 1
                    except Exception as e_img:
                        logging.warning(f"Could not extract/save an image in {file_path.name} (Rel: {rel.rId}): {e_img}")
            logging.info(f"Extracted {len(image_paths)} image(s) from {file_path.name}")

        except Exception as e:
            logging.error(f"Error during basic DOCX processing of {file_path.name}: {e}", exc_info=True)
            results["error"] = f"Failed to fully process base DOCX structure: {e}"

        # 4. Process Extracted Images (OCR and Basic Context)
        stc_pattern = re.compile(r"(?i)\bstc\b")
        for img_ref, img_path in image_paths:
            img_filename = img_path.name
            try:
                ocr_text, diagram_context, graph_data = self._analyze_image(img_path, stc_pattern)
                results["images"][img_filename] = {
                    "ocr_text": ocr_text,
                    "diagram_context": diagram_context,
                    "graph_representation": graph_data,
                    "original_ref": img_ref
                }
            except Exception as e_ocr:
                logging.error(f"Failed OCR/Analysis on {img_filename} from {file_path.name}: {e_ocr}")
                results["images"][img_filename] = {"error": f"Image processing failed: {e_ocr}"}

        return results


    def _process_pdf(self, file_path: Path, temp_dir: Path) -> Dict[str, Any]:
        """
        Processes a .pdf file using PyMuPDF (text, images) and pdfplumber (tables).
        """
        results: Dict[str, Any] = {
            "content_type": "pdf",
            "text_by_page": {},
            "tables": [], # To be filled by pdfplumber
            "images": {}
        }
        image_paths: List[Path] = [] # Store paths from PyMuPDF extraction

        # --- 1. Extract Text and Images using PyMuPDF ('fitz') ---
        doc = None # Initialize doc to None
        try:
            doc = fitz.open(file_path)

            # 1a. Extract Text per Page
            for i, page in enumerate(doc, start=1):
                try:
                    text = page.get_text("text")
                    results["text_by_page"][f"page_{i}"] = text
                except Exception as e_text:
                    logging.warning(f"Could not extract text from page {i} in {file_path.name}: {e_text}")
                    results["text_by_page"][f"page_{i}"] = f"[Text extraction failed on page {i}]"
            logging.info(f"Extracted text from {len(results['text_by_page'])} pages in {file_path.name} using PyMuPDF.")

            # 1b. Extract Images
            image_dir = temp_dir / "images"
            image_dir.mkdir(exist_ok=True)
            img_count = 0
            for i, page in enumerate(doc, start=1):
                try:
                    image_list = page.get_images(full=True)
                    for img_index, img_info in enumerate(image_list, start=1):
                        xref = img_info[0]
                        try:
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            image_ext = base_image["ext"]
                            img_filename = f"page_{i}_img_{img_index}.{image_ext}"
                            img_save_path = image_dir / img_filename
                            with open(img_save_path, "wb") as img_file:
                                img_file.write(image_bytes)
                            image_paths.append(img_save_path)
                            img_count += 1
                        except Exception as e_img_extract:
                            logging.warning(f"Could not extract image xref {xref} on page {i} in {file_path.name}: {e_img_extract}")
                except Exception as e_page_img:
                     logging.warning(f"Could not process images for page {i} in {file_path.name}: {e_page_img}")
            logging.info(f"Extracted {img_count} image file(s) from {file_path.name} using PyMuPDF.")

        except Exception as e:
            logging.error(f"Error during PyMuPDF processing of {file_path.name}: {e}", exc_info=True)
            results["error"] = f"Failed during PyMuPDF processing: {e}"
            # Continue to table extraction if possible, but PyMuPDF part might be incomplete

        finally:
             # Ensure doc is closed if it was opened
            if doc:
                 try:
                      doc.close()
                 except Exception as e_close:
                      logging.warning(f"Minor error closing PyMuPDF document: {e_close}")


        # --- 2. Extract Tables using pdfplumber ---
        pdf_tables = []
        try:
            # Use pdfplumber to open the same PDF file again, specifically for tables
            with pdfplumber.open(file_path) as pdf:
                logging.info(f"Attempting table extraction with pdfplumber for {file_path.name}...")
                total_tables_found = 0
                for i, page in enumerate(pdf.pages):
                    # page.extract_tables() returns a list of tables found on the page.
                    # Each table is a list of lists (rows).
                    # You can pass table_settings configurations if needed.
                    page_tables = page.extract_tables()
                    if page_tables:
                        logging.info(f"pdfplumber found {len(page_tables)} table(s) on page {i+1} of {file_path.name}")
                        for table_data in page_tables:
                            # Clean data: ensure all cells are strings, replace None with ""
                            cleaned_table_data = [
                                [(str(cell) if cell is not None else "") for cell in row]
                                for row in table_data
                                if row is not None # Handle potential None rows if library yields them
                            ]
                            if not cleaned_table_data: continue # Skip empty tables

                            total_tables_found += 1
                            pdf_tables.append({
                                "id": f"table_{total_tables_found}", # Simple incrementing ID
                                "page": i + 1,
                                "data": cleaned_table_data
                                # Note: pdfplumber doesn't provide accuracy/whitespace like Camelot
                            })
                logging.info(f"pdfplumber extracted {total_tables_found} table(s) in total from {file_path.name}")
                results["tables"] = pdf_tables # Assign the extracted tables
        except ImportError:
             # This case should be caught by the top-level check, but included for safety
             logging.error("pdfplumber library not found. Skipping PDF table extraction.")
             results["tables"] = [{"id": "pdfplumber_error", "error": "pdfplumber library not installed."}]
        except Exception as e_pdfplumber:
            logging.error(f"pdfplumber failed to extract tables from {file_path.name}: {e_pdfplumber}", exc_info=True)
            results["tables"].append({"id": "pdfplumber_error", "error": f"pdfplumber processing failed: {e_pdfplumber}"})
            # Keep any tables found before the error occurred, if desired
            # Or clear results["tables"] = [...] if partial results aren't wanted on error

        # --- 3. Process Extracted Images (OCR - from PyMuPDF extraction) ---
        # This part remains the same, processing images extracted by PyMuPDF
        stc_pattern = re.compile(r"(?i)\bstc\b")
        for img_path in image_paths:
            img_filename = img_path.name
            try:
                ocr_text, diagram_context, graph_data = self._analyze_image(img_path, stc_pattern)
                results["images"][img_filename] = {
                    "ocr_text": ocr_text,
                    "diagram_context": diagram_context,
                    "graph_representation": graph_data
                }
            except Exception as e_ocr:
                logging.error(f"Failed OCR/Analysis on {img_filename} from {file_path.name}: {e_ocr}")
                results["images"][img_filename] = {"error": f"Image processing failed: {e_ocr}"}

        return results


    # -------------------------------
    # Helper methods for Image Analysis
    # (These methods remain unchanged)
    # -------------------------------
    def _analyze_image(self, image_path: Path, stc_pattern: re.Pattern) -> Tuple[str, Dict[str, Any], Optional[Dict[str, Any]]]:
        """Performs OCR, extracts basic structure, attempts basic graph extraction."""
        # (Code is identical to previous version)
        ocr_text = ""
        diagram_context = {}
        graph_data = None
        try:
            with Image.open(image_path) as img:
                try:
                    data = pytesseract.image_to_data(img, output_type=Output.DICT)
                    ocr_text = pytesseract.image_to_string(img)
                    logging.info(f"OCR successful for {image_path.name}")
                except pytesseract.TesseractNotFoundError:
                     logging.error("Tesseract executable not found. Cannot perform OCR.")
                     raise RuntimeError("Tesseract not installed or not in PATH.") from None
                except Exception as ocr_err:
                     logging.error(f"Pytesseract OCR failed for {image_path.name}: {ocr_err}")
                     ocr_text = f"[OCR failed: {ocr_err}]"; data = {}

                if stc_pattern.search(ocr_text):
                    ocr_text = stc_pattern.sub("SSS", ocr_text)
                    logging.info(f"Performed STC->SSS replacement in OCR text for {image_path.name}")

                diagram_context = self._extract_basic_diagram_context(data)
                graph = self._create_graph_from_text(ocr_text)
                if graph.nodes:
                    graph_data = nx.node_link_data(graph)
        except FileNotFoundError:
             logging.error(f"Image file not found during analysis: {image_path}"); raise
        except Exception as img_err:
             logging.error(f"Error processing image file {image_path.name}: {img_err}", exc_info=True)
             ocr_text = ocr_text or "[Image analysis failed]"
             diagram_context = diagram_context or {"error": f"Failed to analyze image: {img_err}"}
        return ocr_text, diagram_context, graph_data

    def _extract_basic_diagram_context(self, ocr_data: Dict[str, List]) -> Dict[str, Any]:
        """Extracts text elements and bounding boxes from Tesseract data."""
        # (Code is identical to previous version)
        items = []
        required_keys = ["text", "left", "top", "width", "height", "conf"]
        if not ocr_data or not all(key in ocr_data for key in required_keys):
             return {"recognized_items": [], "notes": "Incomplete OCR data received."}
        num_items = len(ocr_data["text"])
        if not all(len(ocr_data[key]) == num_items for key in required_keys):
            return {"recognized_items": [], "notes": "Inconsistent OCR data lengths."}

        for i in range(num_items):
            confidence = int(float(ocr_data["conf"][i]))
            text = ocr_data["text"][i].strip()
            if text and confidence > 50:
                x, y, w, h = ocr_data["left"][i], ocr_data["top"][i], ocr_data["width"][i], ocr_data["height"][i]
                items.append({"text": text, "bbox": [x, y, x + w, y + h], "confidence": confidence})
        return {
            "recognized_items": items,
            "notes": "Basic context: List of detected text elements with bounding boxes."
        }

    def _create_graph_from_text(self, text: str) -> nx.DiGraph:
        """Creates a directed graph based on simple 'A -> B' patterns found in text."""
        # (Code is identical to previous version)
        G = nx.DiGraph()
        edge_pattern = re.compile(r"""
            \b([\w.-]+)\b \s* (?:-|–|—|=|->|-->|==>|=>) \s* \b([\w.-]+)\b
        """, re.VERBOSE | re.IGNORECASE)
        try:
             edges = edge_pattern.findall(text)
             if edges:
                 logging.info(f"Found {len(edges)} potential graph edges via regex.")
                 for src, dst in edges:
                     src_clean = src.strip('.,:;'); dst_clean = dst.strip('.,:;')
                     if src_clean and dst_clean: G.add_edge(src_clean, dst_clean)
        except Exception as e: logging.error(f"Error during regex-based graph extraction: {e}")
        return G

    # -------------------------------
    # Helper method to open folder
    # -------------------------------
    def _open_folder(self, folder_path: Path) -> None:
        """Opens the specified folder in the system's file explorer."""
        # (Code is identical to previous version)
        if not folder_path.is_dir():
             logging.error(f"Cannot open folder: '{folder_path}' does not exist or is not a directory.")
             return
        try:
            logging.info(f"Attempting to open folder: {folder_path}")
            if sys.platform == "win32": os.startfile(str(folder_path))
            elif sys.platform == "darwin": subprocess.Popen(["open", str(folder_path)])
            else: subprocess.Popen(["xdg-open", str(folder_path)])
        except Exception as e:
            logging.error(f"Could not open folder '{folder_path}': {e}")
            QMessageBox.warning(self, "Open Folder Failed", f"Could not automatically open the folder:\n{folder_path}\n\nReason: {e}")

    def current_time_str(self) -> str:
        """Returns the current time as a formatted string."""
        # Placeholder - requires datetime import if used actively
        # from datetime import datetime
        # return datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        # For now, returning a fixed string as datetime wasn't imported before
        return "Current time check placeholder"


# -------------------------------
# Main entry point
# -------------------------------
if __name__ == "__main__":
    # Ensure Tesseract path is configured if necessary
    # try:
    #     if sys.platform == "win32":
    #         pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    # except Exception as e:
    #      print(f"Warning: Could not set Tesseract path: {e}")

    app = QApplication(sys.argv)
    window = XDoc2JSONApp()
    window.show()
    sys.exit(app.exec_())
